"""
Word文档的结构(包含关系为列表的形式)：
    Document
        Paragraph
            Run
            Run
            Run
            ...
        Paragraph
            Run
            Run
            Run
            ...
        Paragraph
            Run
            Run
            Run
            ...
        ...
            ...
对于paragraph和run对象，使用.text属性获取其文本
docs.save(filename)方法会直接覆盖已有的同名文件
"""
import docx
import docx.enum.text
import docx.shared


def get_structure():
    """观察Document对象的结构"""
    doc = docx.Document('Test\\demo.docx')
    print(len(doc.paragraphs))
    print(type(doc.paragraphs))
    print(doc.paragraphs[0].text)
    print(doc.paragraphs[1].text)
    for i in range(len(doc.paragraphs[1].runs)):
        print(doc.paragraphs[1].runs[i].text)


def get_text():
    """获取文本，忽略格式"""
    doc = docx.Document('Test\\demo.docx')
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    print('\n'.join(full_text))


def change_style():
    doc = docx.Document('Test\\demo.docx')
    print(doc.paragraphs[0].text)
    print(doc.paragraphs[0].style)
    doc.paragraphs[0].sytle = 'Normal'
    print(doc.paragraphs[1].text)
    print(tuple(run.text for run in doc.paragraphs[1].runs))
    doc.paragraphs[1].runs[0] = 'QuoteChar'
    doc.paragraphs[1].runs[1].underline = True
    doc.paragraphs[1].runs[3].underline = True
    doc.save('Test\\restyled.docx')


def write2doc():
    """add_paragraph()和add_run()方法在添加内容后分别返回paragraph对象和run对象，省的多花一步去提取他们"""
    doc = docx.Document('Test\\HelloWorld.docx')
    paragraph1 = doc.add_paragraph('This is the second paragraph added to the file! ')
    paragraph1.add_run('This is being added to the second paragraph !')
    paragraph2 = doc.add_paragraph('This is another paragraph !')
    paragraph2.add_run('This is being added to another paragraph !')
    doc.save('Test\\MultipleParagraph.docx')


def add_title():
    """只有新建的document对象才能使用python-docx模块中内置的style。从已有文档打开的document对象只能使用此文档已有的style"""
    doc = docx.Document()
    doc.add_heading('Heading 0', 0)
    doc.add_heading('Heading 1', 1)
    doc.add_heading('Heading 2', 2)
    doc.add_heading('Heading 3', 3)
    doc.add_heading('Heading 4', 4)
    doc.save('Test\\Heading.docx')


def add_line_page_break():
    doc = docx.Document()
    doc.add_heading('Title_1', 0)
    doc.add_paragraph('This is the first paragraph.')
    paragraph = doc.add_paragraph('This is the second paragraph.')
    paragraph.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)       # 添加换页符。add_break()为Run对象的方法
    doc.add_heading('Title_2', 0)
    doc.add_paragraph('This is the first paragraph after the page break !')
    doc.add_paragraph('This is the second paragraph after the page break !')
    doc.save('Test\\TwoPage.docx')


def add_picture():
    """在文档末尾插入图片，可以同时使用英制和公制单位"""
    doc = docx.Document('Test\\TwoPage.docx')
    doc.add_picture('Test\\sg.png', width=docx.shared.Inches(1.5), height=docx.shared.Cm(4))
    doc.save('Test\\picture.docx')


def new_style():
    """可以新建一个word文档作为模板，在里面设置各种样式。后来的word文档以此为模板打开，可以使用里面的样式，增加内容，然后另存为。"""
    doc = docx.Document('Test\\style.docx')
    doc.paragraphs[0].text = 'This is the first paragraph !'
    doc.paragraphs[0].style = 'Normal'
    paragraph1 = doc.add_paragraph('This is the second paragraph !')
    paragraph1.style = 'mystyle1'
    paragraph2 = doc.add_paragraph('This is the third paragraph !')
    paragraph2.style = 'mystyle1'
    doc.save('Test\\NewStyle.docx')


if __name__ == '__main__':
    new_style()
